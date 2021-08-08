
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfdocument import PDFDocument
from pdfminer.layout import LAParams
from pdfminer.converter import TextConverter
from io import StringIO
from pprint import pprint
import pdfplumber
import textract
import docx
import pathlib
import zipfile

path = r'resumes\Other\non_indian_cvs\EY_Kitman Tsang_Cosec Mgr.docx'
path = r'resumes\best\Arindam_Presales.docx'
path = r'resumes\sample_CVs\Resume_1.docx'
path = r'resumes\Other\non_indian_cvs\DwightIT-QA-Analyst_layout.pdf'
path = r'resumes\sample_CVs\Resume_2.pdf'
path = r'resumes\sample_CVs\my_resume.pdf'
path = r'resumes\Resumes_latest\2MichaelFarros.doc'
path = r'resumes\Resumes_latest\Lawrence Acosta.docx'
path = r'resumes\Resumes_latest\Kevin_Resumev2.docx'


# ------------------------------------
# PDFs - PDFPlumber
# DOCX - Textract (Docx2txt)
# DOC - Textract (antiword)
# ------------------------------------


def extract_text_from_pdf(path):
    output_string = StringIO()
    with open(path, 'rb') as in_file:
        parser = PDFParser(in_file)
        doc = PDFDocument(parser)
        if not doc.is_extractable:
            raise Exception('PDF is not Extractable')
        rsrcmgr = PDFResourceManager()
        device = TextConverter(rsrcmgr, output_string, codec='utf-8', laparams=LAParams())
        interpreter = PDFPageInterpreter(rsrcmgr, device)
        for page in PDFPage.create_pages(doc):
            interpreter.process_page(page)
    return output_string.getvalue()
    # print(output_string.getvalue())

# pdf plumber


def plumb_PDF_text(path):
    text = ''
    hyperlinks = []
    pdf = pdfplumber.open(path)
    for page in pdf.pages:
        # pprint(page.hyperlinks)
        for link in page.hyperlinks:
            hyperlinks.append(link['uri'])
        text = text + '\n' + page.extract_text()
    print(hyperlinks)
    pdf.close()
    return text, hyperlinks


def word_to_Text(path):
    text = textract.process(path)
    return text.decode("utf-8")


def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


def base64ToDocument(path, file_name, extn):
    file_path = pathlib.PurePath(path, file_name + '.' + extn)
    print(file_path)


def unzipFile(path):
    with zipfile.ZipFile(path, 'r') as zip_ref:
        zip_ref.extractall('batch_parsing')
    print('done')


# base64ToDocument('folder1', 'file', 'pdf')
unzipFile(r'uploaded_files\zipped_resume.zip')

# ------------------------------------------------------------------------
# text = extract_text_from_pdf(path)
# print(text)


# print(word_to_Text(path))
# print('------------------------------------------------------------------------')
# print(getText(path))
# # textfromplumber = plumb_PDF_text(path)[0]
# print(textfromplumber)
